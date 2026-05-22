from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.orm import Session
from passlib.context import CryptContext
from database import get_db
from models import User, NprProfile, MentorProfile
import schemas
import io
import urllib.parse
from fastapi.responses import StreamingResponse
from docx import Document
from datetime import datetime
from pydantic import BaseModel
from docx.shared import Pt

class UserStatusUpdate(BaseModel):
    status: str

class UserRoleUpdate(BaseModel):
    role: str

router = APIRouter(prefix="/auth", tags=["Authentication & Profile"])
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

@router.post("/register", response_model=schemas.UserResponse)
def register_user(user: schemas.UserCreate, db: Session = Depends(get_db)):
    db_user = db.query(User).filter(User.email == user.email).first()
    if db_user:
        raise HTTPException(status_code=400, detail="Email уже зарегистрирован")
    
    hashed_pwd = pwd_context.hash(user.password)
    
    if user.email == "admin@sibadi.org":
        new_user = User(email=user.email, hashed_password=hashed_pwd, status="approved", role="admin")
    else:
        new_user = User(email=user.email, hashed_password=hashed_pwd)
    
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    
    if new_user.role == "user":
        new_profile = NprProfile(user_id=new_user.id, experience_years=0, vak_publications=0, rinc_publications=0)
        db.add(new_profile)
        db.commit()
    
    return new_user

@router.post("/login")
def login_user(user: schemas.UserLogin, db: Session = Depends(get_db)):
    db_user = db.query(User).filter(User.email == user.email).first()
    if not db_user or not pwd_context.verify(user.password, db_user.hashed_password):
        raise HTTPException(status_code=401, detail="Неверный email или пароль")
    if db_user.status == "pending":
        raise HTTPException(status_code=403, detail="Ваш аккаунт ожидает подтверждения администратором")
    elif db_user.status == "rejected":
        raise HTTPException(status_code=403, detail="Доступ отклонен администратором")
        
    return {"message": "Успешный вход", "user_id": db_user.id, "role": db_user.role}

@router.put("/profile/{user_id}")
def update_profile(user_id: int, profile_data: schemas.UserProfileUpdate, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.id == user_id).first()
    if not user: raise HTTPException(status_code=404, detail="Пользователь не найден")
    
    if user.role == "user" and user.npr_profile:
        prof = user.npr_profile
        prof.full_name = profile_data.full_name
        prof.birth_date = profile_data.birth_date
        prof.education = profile_data.education
        prof.department = profile_data.department
        prof.position = profile_data.position
        prof.experience_years = profile_data.experience_years
        prof.vak_publications = profile_data.vak_publications
        prof.rinc_publications = profile_data.rinc_publications
        
        calculated_grade = "Не назначен (требования не выполнены)"
        if prof.experience_years >= 3 and prof.rinc_publications >= 5 and prof.vak_publications >= 2:
            calculated_grade = "Грейд 3 (Высокий) - 25 000 руб."
        elif prof.experience_years >= 2 and prof.rinc_publications >= 3 and prof.vak_publications >= 1:
            calculated_grade = "Грейд 2 (Средний) - 20 000 руб."
        elif prof.experience_years >= 1 and prof.rinc_publications >= 1:
            calculated_grade = "Грейд 1 (Начальный) - 15 000 руб."
        
        prof.calculated_grade = calculated_grade
        db.commit()
        return {"message": "Обновлено", "grade_prediction": calculated_grade}
        
    elif user.role == "mentor" and user.mentor_profile:
        prof = user.mentor_profile
        prof.full_name = profile_data.full_name
        db.commit()
        return {"message": "Имя наставника обновлено"}
        
    raise HTTPException(status_code=400, detail="Профиль не доступен для этой роли")

@router.get("/users/pending")
def get_pending_users(db: Session = Depends(get_db)):
    return db.query(User).filter(User.status == "pending").all()

@router.put("/users/{user_id}/status")
def update_user_status(user_id: int, status_data: UserStatusUpdate, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.id == user_id).first()
    if user:
        user.status = status_data.status
        db.commit()
    return {"message": "Статус обновлен"}

@router.get("/profile/{user_id}")
def get_profile(user_id: int, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.id == user_id).first()
    if not user: raise HTTPException(status_code=404, detail="Пользователь не найден")
    
    if user.role == "user" and user.npr_profile:
        return {"profile": user.npr_profile, "grade_prediction": user.npr_profile.calculated_grade}
    elif user.role == "mentor" and user.mentor_profile:
        return {"profile": {"full_name": user.mentor_profile.full_name, "experience_years": 0, "vak_publications": 0, "rinc_publications": 0}}
    return {"profile": {}, "grade_prediction": None}

@router.get("/profile/{user_id}/export")
def export_profile_docx(user_id: int, db: Session = Depends(get_db)):
    profile = db.query(NprProfile).filter(NprProfile.user_id == user_id).first()
    if not profile or not profile.full_name:
        raise HTTPException(status_code=400, detail="Профиль не заполнен")

    grade_text = profile.calculated_grade or "Не назначен"
    
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    for i in range(1, 4):
        if f'Heading {i}' in doc.styles:
            h_style = doc.styles[f'Heading {i}']
            h_style.font.name = 'Times New Roman'
            h_style.font.color.rgb = None

    doc.add_heading('АНКЕТА И ИНДИВИДУАЛЬНЫЙ ПЛАН УЧАСТНИКА ПРОГРАММЫ ПОДДЕРЖКИ МОЛОДЫХ НАУЧНО-ПЕДАГОГИЧЕСКИХ РАБОТНИКОВ', level=1)
    
    doc.add_heading('1. Данные о соискателе', level=2)
    table1 = doc.add_table(rows=11, cols=2)
    table1.style = 'Table Grid'
    
    formatted_birth_date = profile.birth_date or ''
    if profile.birth_date:
        try:
            date_obj = datetime.strptime(profile.birth_date, "%Y-%m-%d")
            formatted_birth_date = date_obj.strftime("%d.%m.%Y")
        except ValueError:
            pass

    data_mapping = [
        ('ФИО', profile.full_name or ''),
        ('Дата рождения', formatted_birth_date),
        ('Образование соискателя (профиль, год выпуска)', profile.education or ''),
        ('Место работы (кафедра)', profile.department or ''),
        ('Должность', profile.position or ''),
        ('Доля ставки', ''), 
        ('Стаж работы в вузе', f"{profile.experience_years} лет"),
        ('Сведения об обучении в аспирантуре (год поступления)', ''),
        ('Место обучения, специальность, научный руководитель', ''),
        ('Условия участия в программе (впервые / год участия)', ''),
        ('Заявляемый уровень грейда на отчетный период', grade_text)
    ]

    for i, (key, val) in enumerate(data_mapping):
        table1.cell(i, 0).text = str(key)
        table1.cell(i, 1).text = str(val)

    doc.add_heading('\n2. Сведения о научных достижениях соискателя', level=2)
    
    doc.add_paragraph('2.1. Перечень опубликованных изданий и научных трудов (за последние 5 лет)')
    table2 = doc.add_table(rows=2, cols=6)
    table2.style = 'Table Grid'
    headers2 = ['№ п/п', 'Наименование', 'Форма', 'Выходные данные', 'Объем', 'Соавторы']
    for i, h in enumerate(headers2): 
        table2.cell(0, i).text = h

    doc.add_paragraph('\n2.2. Участие во всероссийских и международных конференциях (за последние 5 лет)')
    table3 = doc.add_table(rows=2, cols=4)
    table3.style = 'Table Grid'
    headers3 = ['№', 'Наименование конференции', 'Место, даты', 'Форма участия']
    for i, h in enumerate(headers3): 
        table3.cell(0, i).text = h

    doc.add_paragraph('\n2.3. Участие в грантах (за последние 5 лет)')
    table4 = doc.add_table(rows=2, cols=3)
    table4.style = 'Table Grid'
    headers4 = ['№', 'Наименование проекта, сроки', 'Роль в проекте']
    for i, h in enumerate(headers4): 
        table4.cell(0, i).text = h

    doc.add_heading('\n3. Планируемые результаты участия в программе', level=2)
    doc.add_paragraph(f'3.1. Заявляемый уровень грейда: {grade_text}')
    doc.add_paragraph('3.2. Результаты участия в программе:')
    
    table5 = doc.add_table(rows=3, cols=2) 
    table5.style = 'Table Grid'
    table5.cell(0, 0).text = 'Показатель'
    table5.cell(0, 1).text = 'Значение показателя'
    table5.cell(1, 0).text = 'Научная статья, опубликованная в издании ВАК'
    table5.cell(1, 1).text = f"{profile.vak_publications} шт."
    table5.cell(2, 0).text = 'Научная публикация, включенная в РИНЦ'
    table5.cell(2, 1).text = f"{profile.rinc_publications} шт."

    doc.add_paragraph('\n\nУчастник: _________________ / _________________')
    doc.add_paragraph('Наставник: _________________ / _________________')
    doc.add_paragraph('Заведующий кафедрой: _________________ / _________________')

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    filename = f"Anketa_{profile.full_name.replace(' ', '_')}.docx"
    encoded_filename = urllib.parse.quote(filename)
    headers = {'Content-Disposition': f"attachment; filename*=utf-8''{encoded_filename}"}
    return StreamingResponse(file_stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)

@router.get("/users/all")
def get_all_users(db: Session = Depends(get_db)):
    users = db.query(User).all()
    result = []
    for u in users:
        p_data = None
        if u.role == "user" and u.npr_profile: p_data = {"full_name": u.npr_profile.full_name}
        elif u.role == "mentor" and u.mentor_profile: p_data = {"full_name": u.mentor_profile.full_name}
        result.append({"id": u.id, "email": u.email, "status": u.status, "role": u.role, "profile": p_data})
    return result

@router.put("/users/{user_id}/role")
def update_user_role(user_id: int, role_data: UserRoleUpdate, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.id == user_id).first()
    if not user: raise HTTPException(status_code=404, detail="Не найден")
    
    if role_data.role == "mentor" and user.role == "user":
        if user.npr_profile: db.delete(user.npr_profile)
        db.add(MentorProfile(user_id=user.id))
    elif role_data.role == "user" and user.role == "mentor":
        if user.mentor_profile: db.delete(user.mentor_profile)
        db.add(NprProfile(user_id=user.id, experience_years=0, vak_publications=0, rinc_publications=0))
        
    user.role = role_data.role
    db.commit()
    return {"message": f"Роль изменена на {role_data.role}"}

@router.get("/analytics/stats")
def get_analytics_stats(db: Session = Depends(get_db)):
    from models import Ticket
    total_users = db.query(User).filter(User.role == "user").count()
    total_mentors = db.query(User).filter(User.role == "mentor").count()
    
    grades_dist = {"Грейд 1": 0, "Грейд 2": 0, "Грейд 3": 0, "Не назначен": 0}
    all_npr = db.query(NprProfile).all()
    for p in all_npr:
        if p.full_name:
            exp = p.experience_years or 0
            rinc = p.rinc_publications or 0
            vak = p.vak_publications or 0
            if exp >= 3 and rinc >= 5 and vak >= 2: grades_dist["Грейд 3"] += 1
            elif exp >= 2 and rinc >= 3 and vak >= 1: grades_dist["Грейд 2"] += 1
            elif exp >= 1 and rinc >= 1: grades_dist["Грейд 1"] += 1
            else: grades_dist["Не назначен"] += 1
        else:
            grades_dist["Не назначен"] += 1

    return {
        "total_users": total_users,
        "total_mentors": total_mentors,
        "grades": grades_dist,
        "tickets": {
            "total": db.query(Ticket).count(),
            "resolved": db.query(Ticket).filter(Ticket.status == "resolved").count(),
            "pending": db.query(Ticket).filter(Ticket.status == "pending").count()
        }
    }